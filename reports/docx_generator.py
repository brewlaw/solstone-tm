import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

def generate_docx(client_name, attention_name, email, report_date, raw_mark, report_title, page_data, output_filename):
    print("\nDrafting DOCX Legal Opinion Template...")
    
    # --- NEW: Load your pre-made letterhead template! ---
    template_path = os.path.join(os.getcwd(), 'letterhead_template.docx')
    try:
        document = Document(template_path)
    except FileNotFoundError:
        print("🚨 'letterhead_template.docx' not found! Falling back to a blank white document.")
        document = Document()

    # --- PAGE 1: HEADER & METADATA ---
    table = document.add_table(rows=6, cols=2)
    table.autofit = True

    # Helper function to format table rows
    def add_row(row_idx, label, value):
        cell_left = table.cell(row_idx, 0)
        cell_right = table.cell(row_idx, 1)
        cell_left.text = label
        cell_right.text = value
        # Bold the labels
        for paragraph in cell_left.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    add_row(0, "Client Name:\t", client_name)
    add_row(1, "Attention:\t", attention_name)
    add_row(2, "Email:\t", email)
    add_row(3, "Date of Report:\t", report_date)
    add_row(4, "Mark Searched:\t", raw_mark.upper())
    add_row(5, "Type of Search:\t", report_title)

    document.add_paragraph()
    
    # Databases Included Section
    db_para = document.add_paragraph()
    db_para.add_run("Databases Included:\t").bold = True
    
    # Format page numbers dynamically based on what the PDF generated
    def fmt_pgs(start, end):
        return f"(Page {start})" if start == end else f"(Pages {start}-{end})"
        
    db_para.add_run(f"USPTO Trademark Registry {fmt_pgs(page_data['uspto_start'], page_data['uspto_end'])};\n")
    db_para.add_run(f"\t\t\t\tTTB COLA Registry {fmt_pgs(page_data['ttb_start'], page_data['ttb_end'])};\n")
    db_para.add_run(f"\t\t\t\tWeb Search Results {fmt_pgs(page_data['web_start'], page_data['web_end'])}")

    document.add_paragraph()
    document.add_paragraph("Please see a summary of the raw results and a legal opinion regarding third party use of your mark.")

    # --- PAGE BREAK ---
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- PAGE 2: REPORT & ANALYSIS ---
    title_para = document.add_paragraph("TRADEMARK MONITORING REPORT")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.bold = True
        run.underline = True

    document.add_paragraph("Below is a summary of the results we deemed relevant for this report.")
    document.add_paragraph("All of the relevant results are circled on the raw report.")

    # Section Headers for the user to fill in later
    document.add_paragraph("USPTO Trademark Registry").runs[0].bold = True
    document.add_paragraph("[Type relevant USPTO results here...]\n")
    
    document.add_paragraph("TTB COLA Registry").runs[0].bold = True
    document.add_paragraph("[Type relevant TTB results here...]\n")
    
    document.add_paragraph("Web Search Results").runs[0].bold = True
    document.add_paragraph("[Type relevant Web results here...]\n")

    document.add_paragraph("Conclusion:").runs[0].bold = True
    document.add_paragraph("[Type your legal analysis and conclusion here...]\n")

    # --- SIGNATURE BLOCK ---
    # Signature image insertion completely removed!
    document.add_paragraph("By:\tDaniel Christopherson,")
    document.add_paragraph("\tTrademark Attorney")
    document.add_paragraph(f"\t{report_date}")

    # Save Document
    full_path = os.path.join(os.getcwd(), output_filename)
    try:
        document.save(full_path)
        print(f"SUCCESS! DOCX Template saved exactly here:\n -> {full_path}")
    except Exception as e:
        print(f"\n🚨 CRITICAL DOCX ERROR: {e}")