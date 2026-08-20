import streamlit as st
import pandas as pd
import tempfile
import os
import requests
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from dateutil.relativedelta import relativedelta
from playwright.sync_api import sync_playwright
from scrapers.uspto_scraper import scrape_uspto

def calculate_deadline(reg_date_str):
    if not reg_date_str or reg_date_str == "N/A": 
        return "Not Registered"
    try:
        reg_date = pd.to_datetime(reg_date_str)
        current_date = datetime.now()
        
        year_5 = reg_date + relativedelta(years=5)
        year_6 = reg_date + relativedelta(years=6)
        year_9 = reg_date + relativedelta(years=9)
        year_10 = reg_date + relativedelta(years=10)
        
        if current_date <= year_6:
            return f"Sec 8: {year_5.strftime('%Y-%m-%d')} to {year_6.strftime('%Y-%m-%d')}"
        else:
            return f"Sec 8 & 9: {year_9.strftime('%Y-%m-%d')} to {year_10.strftime('%Y-%m-%d')}"
    except:
        return "Unknown"

def run():
    st.header("Trademark Status Report Generator")
    st.write("Generate a clean maintenance and status report for an exact owner/applicant.")

    col1, col2 = st.columns(2)
    with col1:
        owner_name = st.text_input("Exact Owner / Applicant Name", placeholder="e.g. ABC Brewing Co.")
    with col2:
        ic_classes = st.text_input("International Classes (optional, comma-separated)", placeholder="e.g. 032, 033")
        # Adds some invisible spacing so the checkbox aligns nicely with the other column
        st.markdown("<div style='margin-top: 28px;'></div>", unsafe_allow_html=True)
        use_letterhead = st.checkbox("📄 Export Reports on LBL Letterhead", value=False)
        
    exclude_marks = st.text_input("Marks to Exclude (optional, comma-separated)", placeholder="e.g. ABC ALE")

    # Clean, dedicated checkbox placement right above the action button
    use_letterhead = st.checkbox("📄 Export Word Doc on Firm Letterhead", value=False)

    if 'status_report_data' not in st.session_state:
        st.session_state['status_report_data'] = None

    if st.button("Generate Status Report", type="primary"):
        if not owner_name.strip():
            st.warning("Please enter an Owner/Applicant Name.")
            return

        query_parts = [f'ON:"{owner_name.strip().lower()}"', 'LD:true']
        if ic_classes.strip():
            classes = [c.strip().zfill(3) for c in ic_classes.split(",") if c.strip()]
            if classes:
                class_str = " OR ".join([f'IC:{c}' for c in classes])
                query_parts.append(f'({class_str})')
        
        primary_query = " AND ".join(query_parts)
        target_class_list = [c.strip().zfill(3) for c in ic_classes.split(",") if c.strip()] if ic_classes.strip() else None

        with st.spinner("Scraping USPTO trademark records..."):
            try:
                with sync_playwright() as p:
                    browser = p.chromium.launch(
                        headless=True,
                        args=['--no-sandbox', '--disable-dev-shm-usage']
                    )
                    page = browser.new_page()
                    excel_out = f"temp_{owner_name.replace(' ', '_')}.xlsx"
                    raw_results = scrape_uspto(
                        page=page,
                        primary_query=primary_query,
                        excel_filename=excel_out,
                        target_classes=target_class_list
                    )
                    browser.close()
            except Exception as e:
                st.error(f"Error scraping USPTO: {e}")
                return

        if not raw_results:
            st.info("No live marks found matching your query.")
            st.session_state['status_report_data'] = None
            return

        df = pd.DataFrame(raw_results)

        if exclude_marks.strip():
            exclusions = [m.strip().upper() for m in exclude_marks.split(",") if m.strip()]
            df = df[~df['mark'].str.upper().isin(exclusions)]

        df['next_deadline'] = df['reg_date'].apply(calculate_deadline)

        report_df = df[['mark', 'serial', 'reg_number', 'status', 'next_deadline', 'reg_date', 'goods']].copy()
        report_df.columns = ['Mark', 'S/N', 'R/N', 'Status', 'Next Deadline', 'Registration Date', 'Goods & Services']

        # --- 1. PREPARE HTML DATA ---
        html_df = report_df.copy()
        html_df['Mark'] = html_df.apply(
            lambda x: f'<img src="https://tsdr.uspto.gov/img/{x["S/N"]}/large">' if str(x['Mark']).startswith("[Image for ") else x['Mark'],
            axis=1
        )
        
        raw_html_table = html_df.to_html(index=False, escape=False)

        html_report = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 30px; }}
                h1 {{ color: #2F5496; }}
                table {{ border-collapse: collapse; width: 100%; margin-top: 20px; }}
                th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; font-size: 13px; vertical-align: middle; }}
                th {{ background-color: #f2f2f2; color: #333; border-bottom: 2px solid #ddd; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
                img {{ max-height: 80px; max-width: 120px; object-fit: contain; }}
            </style>
        </head>
        <body>
            <h1>Trademark Status Report</h1>
            <p><strong>Owner Searched:</strong> {owner_name}</p>
            <p><strong>Class Filter:</strong> {ic_classes if ic_classes else 'All Live Classes'}</p>
            <p><strong>Date Generated:</strong> {datetime.now().strftime('%B %d, %Y')}</p>
            <hr>
            {raw_html_table.replace('<table border="1" class="dataframe">', '<table>')}
        </body>
        </html>
        """
        
        # --- 2. PREPARE PDF SETUP ---
        class PDF(FPDF):
            def header(self):
                if use_letterhead and os.path.exists("logo.jpg"):
                    self.image("logo.jpg", 10, 8, 30)
                
                self.set_x(45) 
                self.set_font('Arial', 'B', 15)
                self.set_text_color(47, 84, 150)
                self.cell(0, 8, 'Trademark Status Report', 0, 1, 'L')
                
                self.set_x(45)
                self.set_font('Arial', '', 10)
                self.set_text_color(51, 51, 51)
                self.cell(0, 5, f'Owner Searched: {owner_name}', 0, 1, 'L')
                
                self.set_x(45)
                self.cell(0, 5, f'Date Generated: {datetime.now().strftime("%B %d, %Y")}', 0, 1, 'L')
                self.ln(10)

        pdf = PDF(orientation='P')
        pdf.add_page()
        pdf.set_font('Arial', 'B', 8)
        pdf.set_fill_color(242, 242, 242)
        
        headers = ['Mark', 'S/N', 'R/N', 'Status', 'Next Deadline', 'Reg Date', 'Goods']
        col_widths = [35, 17, 17, 13, 35, 18, 55] 
        
        for i in range(len(headers)):
            pdf.cell(col_widths[i], 8, headers[i], 1, 0, 'C', 1)
        pdf.ln()
        
        # --- 3. PREPARE DOCX SETUP ---
        if use_letterhead and os.path.exists("letterhead_template.docx"):
            doc = Document("letterhead_template.docx")
        else:
            doc = Document()
            
        title = doc.add_paragraph()
        run = title.add_run("Trademark Status Report")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(47, 84, 150)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Owner Searched:\t{owner_name}").runs[0].bold = True
        doc.add_paragraph(f"Class Filter:\t{ic_classes if ic_classes else 'All Live Classes'}").runs[0].bold = True
        doc.add_paragraph(f"Date Generated:\t{datetime.now().strftime('%B %d, %Y')}").runs[0].bold = True
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers_docx = ['Mark', 'S/N', 'R/N', 'Status', 'Next Deadline', 'Reg Date', 'Goods']
        for i, h in enumerate(headers_docx):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # --- PROCESS ALL FILES TOGETHER ---
        pdf.set_font('Arial', '', 7)
        line_height = 4
        
        for _, row in report_df.iterrows():
            def clean(val):
                return str(val).replace('“', '"').replace('”', '"').replace("’", "'").encode('latin-1', 'replace').decode('latin-1')
            
            raw_mark = clean(row['Mark'])
            serial_num = clean(row['S/N'])
            
            is_design_mark = False
            img_path = None
            if raw_mark.startswith("[Image for "):
                is_design_mark = True
                try:
                    tsdr_url = f"https://tsdr.uspto.gov/img/{serial_num}/large"
                    response = requests.get(tsdr_url, headers={"User-Agent": "Mozilla/5.0"}, timeout=5)
                    if response.status_code == 200:
                        img_path = f"temp_img_{serial_num}.png"
                        with open(img_path, "wb") as f:
                            f.write(response.content)
                except Exception:
                    pass

            # Update DOCX File
            row_cells = table.add_row().cells
            if is_design_mark and img_path and os.path.exists(img_path):
                p = row_cells[0].paragraphs[0]
                r = p.add_run()
                r.add_picture(img_path, width=Inches(0.8))
            else:
                row_cells[0].text = raw_mark
                
            row_cells[1].text = serial_num
            row_cells[2].text = clean(row['R/N']).replace('nan', '')
            row_cells[3].text = clean(row['Status'])
            row_cells[4].text = clean(row['Next Deadline'])
            row_cells[5].text = clean(row['Registration Date']).replace('nan', '')
            row_cells[6].text = clean(row['Goods & Services'])

            # Update PDF File
            texts = [
                raw_mark, serial_num, clean(row['R/N']).replace('nan', ''), clean(row['Status']),
                clean(row['Next Deadline']), clean(row['Registration Date']).replace('nan', ''), clean(row['Goods & Services'])
            ]
            
            max_lines = 1
            for i, text in enumerate(texts):
                text_width = pdf.get_string_width(text)
                lines = max(1, int((text_width * 1.05) / (col_widths[i] - 2)) + 1)
                if lines > max_lines: max_lines = lines
                    
            row_height = max_lines * line_height
            if is_design_mark and img_path:
                row_height = max(row_height, 20)
            
            if pdf.get_y() + row_height > 275:
                pdf.add_page()
            
            x_start = pdf.get_x()
            y_start = pdf.get_y()
            
            for i, text in enumerate(texts):
                pdf.rect(x_start, y_start, col_widths[i], row_height)
                pdf.set_xy(x_start + 1, y_start + 1)
                if i == 0 and is_design_mark and img_path:
                    try:
                        img_x = x_start + (col_widths[i] - 15) / 2
                        img_y = y_start + (row_height - 15) / 2
                        pdf.image(img_path, x=img_x, y=img_y, w=15, h=15)
                    except Exception:
                        pdf.multi_cell(col_widths[i] - 2, line_height - 0.5, text, border=0, align='C')
                else:
                    align = 'L' if i in [0, 4, 6] else 'C'
                    if not (i == 0 and is_design_mark): 
                        pdf.multi_cell(col_widths[i] - 2, line_height - 0.5, text, border=0, align=align)
                x_start += col_widths[i]
            
            pdf.set_xy(10, y_start + row_height)
            
            if img_path and os.path.exists(img_path):
                os.remove(img_path)

        # Output PDF Bytes
        proper_filename = f"Trademark_Report_{owner_name.replace(' ', '_')}.pdf"
        proper_filepath = os.path.join(tempfile.gettempdir(), proper_filename)
        pdf.output(proper_filepath)
        with open(proper_filepath, "rb") as f:
            pdf_bytes = f.read()
            
        # Output DOCX Bytes
        proper_filename_docx = f"Trademark_Report_{owner_name.replace(' ', '_')}.docx"
        proper_filepath_docx = os.path.join(tempfile.gettempdir(), proper_filename_docx)
        doc.save(proper_filepath_docx)
        with open(proper_filepath_docx, "rb") as f:
            docx_bytes = f.read()

        # Save to state
        st.session_state['status_report_data'] = {
            'owner_name': owner_name,
            'ic_classes': ic_classes,
            'raw_html_table': raw_html_table,
            'html_report': html_report,
            'pdf_bytes': pdf_bytes,
            'docx_bytes': docx_bytes,
            'proper_filepath': proper_filepath,
            'proper_filepath_docx': proper_filepath_docx,
            'proper_filename': proper_filename,
            'proper_filename_docx': proper_filename_docx,
            'count': len(report_df)
        }

    # --- DISPLAY GENERATED REPORT IF PRESENT IN SESSION STATE ---
    if st.session_state.get('status_report_data'):
        data = st.session_state['status_report_data']
        st.success(f"Found {data['count']} mark(s)!")
        
        st.markdown(
            f"""
<style>
    .custom-table {{ border-collapse: collapse; width: 100%; font-size: 14px; margin-bottom: 20px; font-family: sans-serif; }}
    .custom-table th, .custom-table td {{ border: 1px solid #e0e0e0; padding: 12px; text-align: left; vertical-align: middle; }}
    .custom-table th {{ background-color: #f7f7f9; font-weight: 600; color: #31333F; border-bottom: 2px solid #e0e0e0; }}
    .custom-table tr:nth-child(even) {{ background-color: #fbfbfb; }}
    .custom-table img {{ max-height: 70px; max-width: 100px; object-fit: contain; }}
</style>
<div style="overflow-x: auto; border-radius: 8px; border: 1px solid #e0e0e0;">
    {data['raw_html_table'].replace('<table border="1" class="dataframe">', '<table class="custom-table">')}
</div>
            """,
            unsafe_allow_html=True
        )

        dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
        
        with dl_col1:
            st.download_button(
                label="🌐 Download HTML Report",
                data=data['html_report'],
                file_name=f"Trademark_Report_{data['owner_name'].replace(' ', '_')}.html",
                mime="text/html",
                use_container_width=True
            )
            
        with dl_col2:
            st.download_button(
                label="📄 Download Word Doc",
                data=data['docx_bytes'],
                file_name=data['proper_filename_docx'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with dl_col3:
            st.download_button(
                label="📥 Download PDF Report",
                data=data['pdf_bytes'],
                file_name=data['proper_filename'],
                mime="application/pdf",
                use_container_width=True
            )

        with dl_col4:
            if st.button("☁️ Archive to Drive", use_container_width=True):
                from utils.drive_uploader import upload_to_drive
                with st.spinner("Archiving reports..."):
                    with open(data['proper_filepath'], "wb") as f:
                        f.write(data['pdf_bytes'])
                    with open(data['proper_filepath_docx'], "wb") as f:
                        f.write(data['proper_filepath_docx'])
                        
                    pdf_link = upload_to_drive(data['proper_filepath'])
                    docx_link = upload_to_drive(data['proper_filepath_docx'])
                if pdf_link or docx_link:
                    st.success("☁️ Reports archived to Google Drive!")